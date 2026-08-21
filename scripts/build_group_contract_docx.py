# -*- coding: utf-8 -*-
"""Build the Final Project Group Contract .docx that students complete and sign
for Milestone 00.

Output: _final_project/2026F/template/QM474_group_contract.docx

The instruction text of record lives in
_final_project/2026F/milestone_00_meetings_schedule_and_group_contract.md;
this script only renders the fillable form. Re-run after any edit:

    python3 scripts/build_group_contract_docx.py
"""
import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "_final_project/2026F/template/QM474_group_contract.docx")

# ---------- Purdue palette (matches build_schedule_docx.py) ----------
BLACK = "000000"
GOLD = "CFB991"      # Purdue Old Gold
BAND = "F7F5F1"      # subtle warm band
GREY = "6B6B6B"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_borders(cell, color=GOLD, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def para(doc, text="", size=10.5, bold=False, italic=False, color=BLACK,
         space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = RGBColor.from_string(color)
        r.font.name = "Calibri"
    return p


def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{number}. {title}")
    r.font.size = Pt(12)
    r.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor.from_string(BLACK)
    # gold underline rule
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), GOLD)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def prompt(doc, text):
    """Italic grey guidance line under a section heading."""
    para(doc, text, size=9.5, italic=True, color=GREY, space_after=4)


def write_lines(doc, n=3, width_in=6.9):
    """n blank ruled lines for handwriting or typing."""
    t = doc.add_table(rows=n, cols=1)
    t.autofit = False
    for row in t.rows:
        row.height = Inches(0.28)
        c = row.cells[0]
        c.width = Inches(width_in)
        c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        tcPr = c._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "BFBFBF")
        borders.append(bottom)
        tcPr.append(borders)
    return t


def grid(doc, headers, n_rows, widths):
    t = doc.add_table(rows=n_rows + 1, cols=len(headers))
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.width = Inches(widths[j])
        shade(c, GOLD)
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Calibri"
    for i in range(1, n_rows + 1):
        t.rows[i].height = Inches(0.30)
        for j in range(len(headers)):
            c = t.rows[i].cells[j]
            c.width = Inches(widths[j])
            set_cell_borders(c)
            if i % 2 == 0:
                shade(c, BAND)
    return t


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

    # ---------- header ----------
    # Edition-neutral on purpose: the contract is reused each offering, and the
    # folder it lives in already scopes the term. Davi removed the term from the
    # header by hand; keep it out.
    para(doc, "QM 47400 · Predictive Analytics",
         size=10, color=GREY, space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Final Project Group Contract")
    r.font.size = Pt(20)
    r.bold = True
    r.font.name = "Calibri"
    para(doc, "Professor Davi Moreira · Mitch Daniels School of Business",
         size=10, color=GREY, space_after=10)

    para(doc,
         "Complete this contract as a group, agree on every section together, and have all members "
         "sign. Export it as a PDF and submit it with Milestone 00 on Brightspace.",
         size=10, space_after=4)
    para(doc,
         "Fill this in honestly rather than generically. A contract that says \"we will communicate "
         "well\" helps nobody. One that names a channel, a response time, and an escalation point is a "
         "document your group can actually use, and it is what Professor Moreira refers back to if a "
         "contribution dispute is ever raised.",
         size=10, italic=True, color=GREY, space_after=8)

    # ---------- 1. group information ----------
    section_heading(doc, 1, "Group Information")
    grid(doc, ["Group number", "Course section", "Date completed"], 1, [2.3, 2.3, 2.3])
    para(doc, space_after=2)

    # ---------- 2. members ----------
    section_heading(doc, 2, "Members and Contact Information")
    prompt(doc, "One row per member. Use Purdue email addresses.")
    grid(doc, ["Full name", "Purdue email", "Preferred contact method"], 5, [2.1, 2.6, 2.2])
    para(doc, space_after=2)

    # ---------- 3. communication ----------
    section_heading(doc, 3, "Communication Norms")
    prompt(doc, "Name your primary channel and the response time every member commits to. "
                "Example: \"Slack, replies within 24 hours on weekdays.\"")
    write_lines(doc, 3)
    para(doc, space_after=2)

    # ---------- 4. meetings ----------
    section_heading(doc, 4, "Meeting Cadence")
    prompt(doc, "When does your group meet, where, and who is responsible for scheduling it?")
    write_lines(doc, 3)
    para(doc, space_after=2)

    # ---------- 5. roles ----------
    section_heading(doc, 5, "Roles and Responsibilities")
    prompt(doc, "Who takes point on each area? Roles may rotate. If they do, say how and when.")
    t = grid(doc, ["Area", "Who takes point", "Notes"], 6, [1.9, 2.2, 2.8])
    for i, area in enumerate(["Data collection and cleaning", "Modeling and evaluation",
                              "Writing and abstract", "Poster design",
                              "Brightspace submissions", "Other"], start=1):
        c = t.rows[i].cells[0]
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(area)
        r.font.size = Pt(9.5)
        r.font.name = "Calibri"
    para(doc, space_after=2)

    doc.add_page_break()

    # ---------- 6. decisions ----------
    section_heading(doc, 6, "Decision Making")
    prompt(doc, "How does your group settle a disagreement about project direction? "
                "Majority vote, consensus, or the member who owns that area decides?")
    write_lines(doc, 3)
    para(doc, space_after=2)

    # ---------- 7. work distribution ----------
    section_heading(doc, 7, "Work Distribution and Internal Deadlines")
    prompt(doc, "How do tasks get assigned, where are they tracked, and how far ahead of a "
                "Brightspace deadline does your group set its own internal deadline?")
    write_lines(doc, 4)
    para(doc, space_after=2)

    # ---------- 8. accountability ----------
    section_heading(doc, 8, "Accountability")
    prompt(doc, "This is the section groups most want to leave vague and the one that matters most "
                "if something goes wrong. Be specific.")
    para(doc, "What does your group do when a member misses an internal deadline?",
         size=10, bold=True, space_after=2, space_before=4)
    write_lines(doc, 2)
    para(doc, "What does your group do when a member is unreachable? "
              "After how many days does that trigger a response?",
         size=10, bold=True, space_after=2, space_before=6)
    write_lines(doc, 2)
    para(doc, "At what point does your group bring the issue to Professor Moreira?",
         size=10, bold=True, space_after=2, space_before=6)
    write_lines(doc, 2)
    para(doc, space_after=2)

    # ---------- 9. acknowledgement + signatures ----------
    section_heading(doc, 9, "Acknowledgement and Signatures")
    para(doc,
         "By signing below, each member confirms that they helped write this contract, agree to its "
         "terms, and understand that intra-group peer evaluation is worth 20% of the Final Project "
         "grade and is assessed per member.",
         size=10, space_after=6)
    grid(doc, ["Full name", "Signature", "Date"], 5, [2.4, 3.0, 1.5])

    para(doc, space_after=4)
    para(doc,
         "Submit as NN_group_contract.pdf, where NN is your group number, "
         "with the Milestone 00 assignment on Brightspace.",
         size=9.5, italic=True, color=GREY, space_before=10,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
