# -*- coding: utf-8 -*-
"""Rebuild 2026F_predictive_analytics_QM474_schedule.docx from the live
webpage schedule table in schedule.qmd (single source of truth)."""
import io, os, re
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Emu

REPO = "/Users/dcordeir/Dropbox/academic/cursos/cursos-davi/predictive_analytics/2026F_predictive_analytics_QM474"
SRC  = os.path.join(REPO, "schedule.qmd")
OUT  = os.path.join(REPO, "_syllabus/2026F/2026F_predictive_analytics_QM474_schedule.docx")

# ---------- Purdue palette -------------------------------------------------
BLACK   = "000000"
GOLD    = "CFB991"   # Purdue Old Gold
BAND    = "F7F5F1"   # subtle warm band
NOCLASS = "EFEFEF"
ACCENT  = "EFE7D6"   # gold tint for the three marquee events
GREY    = RGBColor(0x76, 0x76, 0x76)
# Calibri/Cambria is the theme of the official syllabus docx, so the schedule
# ships as a visual set with it (and Calibri is present in every Office install).
FONT    = "Calibri"
HEAD_FONT = "Cambria"

# ---------- helpers --------------------------------------------------------
def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

def set_cell_borders(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge, spec in kw.items():
        e = borders.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge); borders.append(e)
        for k, v in spec.items():
            e.set(qn('w:' + k), str(v))

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

def set_table_cell_margins(table, top=60, start=90, bottom=60, end=90):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for name, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        n = OxmlElement('w:' + name)
        n.set(qn('w:w'), str(val)); n.set(qn('w:type'), 'dxa')
        mar.append(n)
    tblPr.append(mar)

def style_run(run, size=9, bold=False, italic=False, color=None, mono=False):
    run.font.name = "Consolas" if mono else FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # east-asian + cs font so Word does not substitute
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(a), "Consolas" if mono else FONT)
    return run

def add_hyperlink(par, url, text, size=8, bold=False):
    """Real external hyperlink.

    The run is created through python-docx's own API and only then moved inside
    the w:hyperlink wrapper. Hand-building w:rPr risks an out-of-order child
    sequence (ECMA-376 CT_RPr is an ordered sequence: rFonts, b, color, sz, u),
    which makes Word prompt to repair the file.
    """
    r_id = par.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    run = par.add_run(text)
    style_run(run, size, bold=bold, color=RGBColor(0x1F, 0x4E, 0x79))
    run.font.underline = True
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), r_id)
    par._p.remove(run._element)
    link.append(run._element)
    par._p.append(link)


INLINE = re.compile(r'(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)')

def write_inline(par, text, size=9, base_bold=False, base_italic=False, color=None):
    """Render a markdown-inline string into runs on `par`."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith('**') and piece.endswith('**'):
            style_run(par.add_run(piece[2:-2]), size, True, base_italic, color)
        elif piece.startswith('`') and piece.endswith('`'):
            style_run(par.add_run(piece[1:-1]), size - 0.5, base_bold, base_italic, color, mono=True)
        elif piece.startswith('*') and piece.endswith('*'):
            style_run(par.add_run(piece[1:-1]), size, base_bold, True, color)
        else:
            style_run(par.add_run(piece), size, base_bold, base_italic, color)

def tune(par, space_before=1, space_after=1, align=None):
    pf = par.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if align is not None: pf.alignment = align
    return par

# ---------- parse the webpage schedule table -------------------------------
lines = io.open(SRC, encoding="utf-8").read().split("\n")
start = lines.index("::: overflow-table")
end   = start + 1 + lines[start + 1:].index(":::")
rows = []
for ln in lines[start + 1:end]:
    if not ln.startswith("|"):
        continue
    cells = [c.strip() for c in ln.rstrip()[1:-1].split("|")]
    if set("".join(cells)) <= set("-: "):      # separator row
        continue
    assert len(cells) == 5, cells
    rows.append(cells)
header, data = rows[0], rows[1:]
assert header == ['Wk', 'Date', 'Topic', 'Notebook', 'Assessment / Notes'], header
print("parsed %d session rows from schedule.qmd" % len(data))

NB_RE   = re.compile(r'\*\*(nb\d+)\*\*')
# NOTE: the badge markdown holds TWO colab URLs — the badge .svg image and the
# real notebook target. Anchor on /github/ so we never link the image asset.
LINK_RE = re.compile(r'\]\((https://colab\.research\.google\.com/github/[^)]+)\)')

# ---------- milestones (carried forward; matches _final_project/2026F) -----
MILESTONES = [
    ("M00", "Group Contact Confirmation",                  "Sun Sep 6"),
    ("M01", "Initial Project Proposal",                    "Sun Sep 20"),
    ("M02", "Expanded Project Outline",                    "Sun Sep 27"),
    ("M03", "Project Draft Abstract",                      "Sun Oct 4"),
    ("M05", "Applying to the Conference",                  "Sun Oct 11"),
    ("M06", "Simple Model & Performance Evaluation",       "Sun Oct 18"),
    ("M08", "More Complex Models & Performance Evaluation","Sun Oct 25"),
    ("M09", "Poster First Draft",                          "Sun Nov 1"),
    ("M10", "Final Poster Submission (NN.pdf)",            "Sun Nov 8"),
    ("M11", "Poster Presentation Planning",                "Sun Nov 15"),
    ("M12", "LinkedIn Post Invitation",                    "Sun Nov 15"),
    ("—",   "URC Poster Presentation (required)",          "Tue Nov 17"),
    ("—",   "Intra-group Peer Evaluation",                 "Fri Dec 11"),
]

# ---------- build ----------------------------------------------------------
doc = Document()

st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11), Inches(8.5)
sec.left_margin = sec.right_margin = Inches(0.5)
sec.top_margin = Inches(0.45); sec.bottom_margin = Inches(0.45)
USABLE = Inches(10.0)

# --- title block
p = tune(doc.add_paragraph(), 0, 0)
style_run(p.add_run("QM 47400 · Predictive Analytics"), 20, bold=True)
p = tune(doc.add_paragraph(), 1, 0)
style_run(p.add_run("Course Schedule · Fall 2026"), 13, bold=True, color=RGBColor(0x8E, 0x76, 0x4C))
p = tune(doc.add_paragraph(), 2, 8)
style_run(p.add_run(
    "Mitch Daniels School of Business, Purdue University   ·   "
    "Professor Davi Moreira   ·   Sections 001 & 002, Monday / Wednesday / Friday, WTHR 114   ·   "
    "August 24 – December 11, 2026"), 9, color=GREY)

# --- main schedule table
WIDTHS = [0.05, 0.105, 0.42, 0.125, 0.30]     # sums to 1.0
assert abs(sum(WIDTHS) - 1.0) < 1e-9
tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
tbl.autofit = False
tbl.allow_autofit = False
set_table_cell_margins(tbl)

def set_widths(table):
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = Emu(int(USABLE * table._widths[i]))

tbl._widths = WIDTHS

hdr = tbl.rows[0]
repeat_header(hdr); cant_split(hdr)
for i, label in enumerate(header):
    cell = hdr.cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade(cell, BLACK)
    par = tune(cell.paragraphs[0], 3, 3,
               WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT)
    style_run(par.add_run(label), 9.5, bold=True, color=RGBColor(0xCF, 0xB9, 0x91))

prev_wk = None
for cells in data:
    wk, date, topic, nb, note = cells
    row = tbl.add_row(); cant_split(row)

    is_noclass  = topic.lstrip().startswith('*No class')
    is_marquee  = ('MIDTERM EXAM' in topic or 'Undergraduate Research Conference' in topic
                   or 'SHOWCASE' in topic)
    band = int(wk.strip('* ')) % 2 == 0
    fill = ACCENT if is_marquee else (NOCLASS if is_noclass else (BAND if band else None))
    new_week = wk != prev_wk

    for i, raw in enumerate(cells):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if fill: shade(cell, fill)
        if new_week:      # stronger rule at each week boundary
            set_cell_borders(cell, top={'val': 'single', 'sz': 12, 'color': 'BFB39A'})
        par = tune(cell.paragraphs[0], 2, 2,
                   WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT)

        if i == 3:                                   # Notebook column
            m, link = NB_RE.search(raw), LINK_RE.search(raw)
            if m:
                style_run(par.add_run(m.group(1)), 9.5, bold=True)
                par.add_run().add_break()
                add_hyperlink(par, link.group(1), "Open in Colab", size=8)
            else:
                style_run(par.add_run("—"), 9, color=GREY)
        elif is_noclass and i == 2:
            write_inline(par, raw, 9, color=GREY)
        else:
            write_inline(par, raw, 9,
                         base_bold=(is_marquee and i == 1))
    prev_wk = wk

set_widths(tbl)

# --- milestones
p = tune(doc.add_paragraph(), 14, 2)
style_run(p.add_run("Final Project Milestones"), 13, bold=True)
p = tune(doc.add_paragraph(), 0, 6)
style_run(p.add_run(
    "Group capstone. Deliverables are due Sundays, 11:59 PM, except the final poster, "
    "pinned to Sun Nov 8 — nine days before the conference.  Milestone numbering follows the "
    "2026F reference documents; M04 / M07 / M13 are intentionally unused."), 9, color=GREY)

m_tbl = doc.add_table(rows=1, cols=3)
m_tbl.style = 'Table Grid'
m_tbl.autofit = False; m_tbl.allow_autofit = False
m_tbl._widths = [0.08, 0.62, 0.30]
set_table_cell_margins(m_tbl)
mh = m_tbl.rows[0]; repeat_header(mh); cant_split(mh)
for i, label in enumerate(("#", "Milestone", "Due")):
    cell = mh.cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade(cell, BLACK)
    par = tune(cell.paragraphs[0], 3, 3,
               WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    style_run(par.add_run(label), 9.5, bold=True, color=RGBColor(0xCF, 0xB9, 0x91))

for idx, (num, name, due) in enumerate(MILESTONES):
    row = m_tbl.add_row(); cant_split(row)
    key = num == "—"
    for i, raw in enumerate((num, name, due)):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if key: shade(cell, ACCENT)
        elif idx % 2 == 0: shade(cell, BAND)
        par = tune(cell.paragraphs[0], 2, 2,
                   WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
        style_run(par.add_run(raw), 9, bold=key)
set_widths(m_tbl)

p = tune(doc.add_paragraph(), 10, 0)
style_run(p.add_run(
    "Subject to change. While I will try to adhere to the course schedule as much as possible, "
    "I also want to adapt to your learning pace and style; the syllabus and course plan may change "
    "during the term. The official source of record is the course Brightspace page."), 8.5,
    italic=True, color=GREY)

# --- footer with page numbers
footer_par = sec.footer.paragraphs[0]
footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer_par.add_run("QM 47400 · Predictive Analytics · Fall 2026 · Page "), 8, color=GREY)
for instr in ('PAGE',):
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), instr)
    footer_par._p.append(fld)
style_run(footer_par.add_run(" of "), 8, color=GREY)
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'NUMPAGES')
footer_par._p.append(fld)

doc.save(OUT)
print("WROTE", OUT)
print("sessions: %d   milestones: %d" % (len(data), len(MILESTONES)))
