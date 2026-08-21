# -*- coding: utf-8 -*-
"""Build the Mitch Daniels School of Business HONORS SYLLABUS that an honors
contract requires as its "modified syllabus" attachment.

Output: _syllabus/2026F/honors_contract/QM47400_2026F_honors_syllabus.docx

Structure follows the DSB template of record (OBJECTIVES / DELIVERABLES /
DEADLINES / GRADING SCHEME / ADDITIONAL INFORMATION). The policy text of record
is the "Honors Contract (Optional)" section of syllabus.qmd; this script only
renders it into the attachable form. Re-run after any edit:

    python3 scripts/build_honors_syllabus_docx.py

Fill in the student's name and PUID in the header line before sending; the rest
is the same for every contract this term, including a group contract.
"""
import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "_syllabus/2026F/honors_contract/QM47400_2026F_honors_syllabus.docx")

BLACK = "000000"
GOLD = "CFB991"      # Purdue Old Gold
BAND = "F7F5F1"
GREY = "6B6B6B"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_borders(cell, color=GOLD, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def para(doc, text="", size=10.5, bold=False, italic=False, color=BLACK,
         space_after=6, space_before=0, align=None, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = RGBColor.from_string(color)
        r.font.name = "Calibri"
    return p


def rich(doc, chunks, size=10.5, space_after=6, indent=0.0):
    """chunks = [(text, bold), ...] rendered as one paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold in chunks:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor.from_string(BLACK)
    return p


def section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.size = Pt(11.5)
    r.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor.from_string(BLACK)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), GOLD)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def table(doc, headers, rows, widths, header_size=9.5, body_size=9.5):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.width = Inches(widths[j])
        shade(c, GOLD)
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(header_size)
        r.font.name = "Calibri"
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.width = Inches(widths[j])
            set_cell_borders(c, color="D9D9D9")
            if i % 2 == 0:
                shade(c, BAND)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            bold = val.startswith("**") and val.endswith("**")
            r = p.add_run(val.strip("*"))
            r.bold = bold
            r.font.size = Pt(body_size)
            r.font.name = "Calibri"
    return t


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)

    para(doc, "MITCH DANIELS SCHOOL OF BUSINESS HONORS PROGRAM", size=9,
         bold=True, color=GREY, space_after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    para(doc, "MITCH DANIELS SCHOOL OF BUSINESS HONORS SYLLABUS", size=15,
         bold=True, space_after=8)

    rich(doc, [("COURSE: ", True), ("QM 47400, Predictive Analytics", False)], space_after=2)
    rich(doc, [("ACADEMIC PERIOD: ", True), ("Fall 2026 (August 24 - December 11, 2026)", False)], space_after=2)
    rich(doc, [("INSTRUCTOR: ", True), ("Professor Davi Moreira, dcordeir@purdue.edu", False)], space_after=2)
    rich(doc, [("STUDENT: ", True), ("_______________________________     PUID: ______________", False)], space_after=2)
    para(doc, "This honors syllabus accompanies the standard QM 47400 syllabus. Everything in the "
              "standard syllabus applies unchanged except the grading scheme restated below.",
         size=9.5, italic=True, color=GREY, space_after=4)

    # ---------------- OBJECTIVES ----------------
    section_heading(doc, "HONORS CONTRACT OBJECTIVES")
    para(doc,
         "The objective of this honors contract is for the student to own an individual line of "
         "research inside the course's team-based predictive analytics project, and to carry it "
         "through to a public presentation. The standard course asks each team to build, validate, "
         "and communicate a predictive model for a business stakeholder. The honors version asks "
         "the student, individually, to formulate a scholarly question of their own within that "
         "problem, to answer it with a modeling approach that goes beyond the team's scope, to "
         "evaluate it under the course's cross-validation protocol against the team's own baseline "
         "on identical folds, and to present the result under their own name at the Purdue "
         "Undergraduate Research Conference on November 17, 2026.")
    para(doc,
         "The honors work is therefore additional and individual, while remaining anchored in the "
         "course the student is already taking: it adds an independent question, an independent "
         "analysis, and independently attributed authorship on top of collaborative group work.")

    # ---------------- DELIVERABLES ----------------
    section_heading(doc, "HONORS CONTRACT DELIVERABLES")
    para(doc, "The student submits three individual deliverables through Brightspace.", space_after=6)

    rich(doc, [("1. Honors Question Memo (about 2 pages). ", True),
               ("A written statement of the specific scholarly question the student will own inside "
                "the team's project: the question itself, why it matters to the stakeholder, three "
                "to five sources the work builds on, and the method by which the question will be "
                "answered.", False)], indent=0.15)
    rich(doc, [("2. Honors Analysis (Colab notebook plus a 3-page results memo). ", True),
               ("One modeling extension beyond the set of models the team develops - for example an "
                "alternative algorithm, an alternative target or cost structure, a fairness or "
                "sensitivity analysis, or an alternative feature-engineering strategy. Performance "
                "is estimated with k-fold cross-validation on the training data and compared with "
                "the team's baseline through paired per-fold differences on identical folds, "
                "against a margin declared in advance. The memo states the question, the method, "
                "the comparison, the limitations, and the recommendation to the stakeholder.", False)],
         indent=0.15)
    rich(doc, [("3. Honors Section and Conference Walkthrough. ", True),
               ("One clearly labeled section on the team's conference poster presenting the "
                "student's extension and carrying the student's name, plus a five-minute individual "
                "walkthrough of that section delivered at the Purdue Undergraduate Research "
                "Conference. Where the entire team holds honors contracts, the poster as a whole is "
                "the honors deliverable and no separate section is required.", False)], indent=0.15)
    para(doc,
         "In addition, the student meets individually with the instructor three times during the "
         "term, for approximately fifteen minutes each, in late September, late October, and after "
         "the conference, to review progress on the deliverables above.", space_before=4)

    # ---------------- DEADLINES ----------------
    section_heading(doc, "HONORS CONTRACT DEADLINES")
    para(doc, "All deliverables are due at 11:59 p.m. on the dates below, submitted through Brightspace. "
              "The honors contract requires specific deadlines, so they are stated here even though the "
              "course syllabus defers deliverable dates to Brightspace.",
         size=9.5, italic=True, color=GREY, space_after=5)
    table(doc,
          ["Deliverable", "Due", "Aligned course milestone"],
          [["**Honors Question Memo**", "Sunday, September 20, 2026", "M01 - Initial Project Proposal"],
           ["**Honors Analysis**", "Sunday, October 25, 2026", "M08 - More Complex Models and Performance Evaluation"],
           ["**Honors Section (in the team poster)**", "Tuesday, November 10, 2026", "M10 - Final Poster Submission"],
           ["**Conference Walkthrough**", "Tuesday, November 17, 2026", "Purdue Undergraduate Research Conference"]],
          widths=[2.5, 1.9, 2.6])

    # ---------------- GRADING ----------------
    section_heading(doc, "HONORS GRADING SCHEME")
    para(doc,
         "Ten percent of the course grade is reallocated from the group Final Project to the "
         "individual Honors Research Extension, giving the honors component a noticeable impact on "
         "the final grade, unless the entire team holds honors contracts. The Final Project retains "
         "all five of its standard components; each is scaled proportionally from 35% to 25% of the "
         "course grade. Every other assessment is unchanged.", space_after=6)
    table(doc,
          ["Assessment", "Standard", "Honors"],
          [["Attendance", "1%", "1%"],
           ["Participation", "4%", "4%"],
           ["Quizzes", "15%", "15%"],
           ["Midterm Exam", "20%", "20%"],
           ["Course Case Competition (Kaggle)", "20%", "20%"],
           ["Final Project", "35%", "25%"],
           ["Poster-to-Product", "5%", "5%"],
           ["**Honors Research Extension**", "-", "**10%**"],
           ["**Total**", "**100%**", "**100%**"]],
          widths=[4.0, 1.5, 1.5])
    para(doc, "The Honors Research Extension is graded out of the 10 points as follows: Honors "
              "Question Memo 2 points, Honors Analysis 5 points, Honors Section and Conference "
              "Walkthrough 3 points. Each is assessed on clarity of the question, soundness of the "
              "evaluation, and quality of the communication to a non-technical stakeholder.",
         size=9.5, space_before=5)
    para(doc, "The honors section on the poster is graded exclusively as part of the Honors Research "
              "Extension. It is not evaluated under the group poster rubric and therefore cannot "
              "raise or lower the team's poster grade.", size=9.5, space_before=4)
    para(doc, "During the term the Brightspace gradebook reflects only the grade earned through the "
              "regular course requirements. After the student completes the full honors contract, "
              "the instructor manually adjusts the final course grade to incorporate the Honors "
              "Research Extension.", size=9.5, space_before=4)

    # ---------------- ADDITIONAL ----------------
    section_heading(doc, "ADDITIONAL INFORMATION")
    para(doc,
         "This contract is designed so that it may also serve as the student's John Martinson "
         "Honors College Scholarly Project. It produces new knowledge that is individually "
         "attributable - the question, the analysis, and the labeled poster section are the student's "
         "own - and it is presented publicly at the Purdue Undergraduate Research Conference. The "
         "instructor serves as faculty mentor. Approval of the Scholarly Project remains a decision "
         "of the Honors College: the student submits the proposal through the Honors College portal "
         "by the fall deadline of October 1, 2026, and files completion verification through the "
         "same portal afterward. The Honors Question Memo is written to double as the proposal draft.")
    para(doc,
         "A group honors contract is available when more than one student on the same team "
         "contracts the course, limited to the course maximum group size. Each student still owns a "
         "distinct question and submits their own memo and analysis; only the administration is "
         "shared. Where every member of a team holds a contract, the team's poster as a whole is "
         "the honors deliverable.")

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
